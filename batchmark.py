#!/usr/bin/env python2
# vim: et:ts=4:textwidth=80

# Automark
#
# David Llewellyn-Jones
# Liverpool John Moores University
# 18/12/2014
# Released under the GPL v.3. See the LICENSE file for more details.

"""
Check multiple programs against task requirements

This script allows multiple programs to be checked either locally or
using the ideone api.

David Llewellyn-Jones, Posco Tso
Liverpool John Moores University
18/12/2014
Released under the GPL v.3. See the LICENSE file for more details.

"""

import os

import automark
import automarktask1
import automarktask2
import automarktask3

from argparse import ArgumentParser
from docx import Document
from time import time
from xlrd import open_workbook
from zipfile import ZipFile

__all__ = ('BatchMark')


# Task number
# Folder containing students' folders
# Marker's initials (optional)
# Temp build folder (optional)
# Feedback template (optional)
# Student details list (optional)
# Summary output (optional)

class BatchMark(object):
    """
    Mark multiple code submissions and generate feedback sheets.
    
    Initialise with the various attributes outlined below. Then call the
    go() method to start the batch marking process.

    Attributes:
        task: Task number
        marking_dir: Folder containing students' folders
        marker_name: Marker's initials (optional)
        build_dir: Temp build folder (optional)
        feedback_doc_name: Feedback template (optional)
        marking_sheet_name: Student details list (optional)
        summary-out: Summary output (optional)

    """

    def __init__(
            self, task, marking_dir, marker_name, build_dir, 
            feedback_doc_name, marking_sheet_name, summary_out):
        """
        Initialise the BatchMark class with the attributes needed for it
        to run, and print out those attributes.
        """
        print ('Task: ' + str(task))
        print ("Folder to check in: " + marking_dir)
        print ('Marker name: ' + marker_name)
        print ("Build output folder: " + build_dir)
        print ("Feedback template: " + feedback_doc_name)
        print ("Marking sheet: " + marking_sheet_name)
        print ("Summary output: " + summary_out)

        self._task = task
        self._marking_dir = marking_dir
        self._marker_name = marker_name
        self._buid_dir = build_dir
        self._feedback_doc_name = feedback_doc_name
        self._marking_sheet_name = marking_sheet_name
        self._summary_out = summary_out

        # Select the appropriate task
        tasks = [automark, automarktask1, automarktask2, automarktask3]
        if (self._task < len(tasks)) and (self._task >= 0):
            self._task_specific = tasks[self._task]
        else:
            self._task_specific = tasks[0]
            print 'Task number out of bounds. Applying general checks.'

        #load student feedback form as a template
        self._feedback_document = Document(feedback_doc_name)
        #load my marking sheet 'PT' from workbook
        self._marking_sheet = open_workbook(
            marking_sheet_name).sheet_by_name(marker_name)

    #do things
    def go(self):
        """
        Start the marking process based on the initialisation parameters.
        """
        #username to firstname lastname map/dictionary
        self._name_map = {}
        self._construct_name_map()
        with open(self._summary_out, 'w') as self._summary:
            self._output_csv_co(
                ['Username', 'Name'])
            self._output_csv_co(
                self._task_specific.Automark.get_scores_structure())
            self._output_csv_co(
                self._task_specific.Automark.get_internal_stats_structure())
            self._output_csv_nl(
                self._task_specific.Automark.get_output_checks_structure())
            self._create_new_feedback_document()

    def _unzip_submission(self, student_dir):
        """
        Internal method, cycle through the folders in the given path and
        extract the contents of any zip archives found.
        """
        archive_file = ''
        for check_dir, _, file in os.walk(student_dir):
            if '__MACOSX' not in check_dir.split('/'):
                for file_name in file:
                    if file_name.endswith('.zip'):
                        archive_file = os.path.join(check_dir, file_name)
        if archive_file != '':
            with ZipFile(archive_file ,'r') as archive:
                archive.extractall(student_dir)
                archive.close()
        else:
            print ('No zip archive')

    def _output_csv(self, items):
        """
        Internal method, output the item list to a csv file. Items are 
        separated by commas, but there's no comma at the start or end.
        """
        count = 0
        for item in items:
            self._summary.write('\"')
            self._summary.write(str(item))
            self._summary.write('\"')
            if count < (len(items) - 1):
                self._summary.write(', ')
            count += 1

    def _output_csv_co(self, items):
        """
        Internal method, output the item list to a csv file. Items are 
        separated by commas and a comma is acced to the end for 
        continuation.
        """
        self._output_csv(items)
        self._summary.write(', ')

    def _output_csv_nl(self, items):
        """
        Internal method, output the item list to a csv file. Items are
        separated by commas and a newline is output at the end.
        """
        self._output_csv(items)
        self._summary.write('\n')

    def _mark (self, directory):
        """
        Internal method, return the last java file found in the 
        directory.
        """
        java_file = ''
        for check_dir, _, file in os.walk(directory):
            if '__MACOSX' not in check_dir.split('/'):
                for file_name in file:
                    if file_name.endswith('.java'):
                        java_file = os.path.join(check_dir, file_name)
        return java_file

    def _create_new_feedback_document(self):
        """
        Internal method, cycle through the folders, unzip any archives, 
        compile and execute the java programs, test the results, output 
        a completed feedback file, write out the results to the csv 
        summary file.
        """
        print
        for student_dir, _, file in os.walk(self._marking_dir):
            student_dir_name = os.path.relpath(student_dir, self._marking_dir)

            #print student_dir
            if (student_dir_name is not '.') and (
                    student_dir_name in self._name_map):
                print 'Student: {}'.format(student_dir_name)
                student_name = self._name_map[student_dir_name][0] \
                    + ' ' + self._name_map[student_dir_name][1]
                self._unzip_submission(student_dir)
                java_path = self._mark(student_dir)
                if java_path == '':
                    print 'No java file'
                    self._write_student_name_to_document(
                        student_dir, student_dir_name, student_name)
                else:
                    #print 'file: {}'.format(java_path)
                    marks = self._task_specific.Automark(
                        java_path, 'credentials.txt', self._buid_dir)
                    self._write_details_to_document(
                        student_dir, student_dir_name, student_name, marks)
                    self._write_comments_to_document(
                        student_dir, student_dir_name, marks)
                    self._output_csv_co(
                        [student_dir_name, student_name])
                    self._output_csv_co(
                        marks.get_scores())
                    self._output_csv_co(
                        marks.get_internal_stats())
                    self._output_csv_nl(
                        marks.get_output_checks())

                #just do something extra
                #self._unzip_submission(student_dir)

    def _write_details_to_document(
            self, student_dir, student_dir_name, student_name, marks):
        """
        Internal method, write out student details and the results of the 
        marking to a feedback document.
        """
        #default cell for student's firstname lastname
        filename = self._feedback_doc_name.replace(
            'username', student_dir_name)
        self._feedback_document.paragraphs[0].text = \
            self._feedback_document.paragraphs[0].text.replace(
            '<task>', str(self._task))
        self._feedback_document.tables[0].cell(1,0).text = student_name
        self._feedback_document.tables[0].cell(1,1).text = student_dir_name
        #self._feedback_document.tables[0].cell(1,2).text = self._marker_name
        self._feedback_document.tables[0].cell(1,2).text = 'auto'
        #print student_dir+'/'+filename
        # Clear tables
        for row in range(1,16):
            self._feedback_document.tables[2].cell(row, 2).text = ''
    
        execution_score = marks.get_execution_score()
        indentation_score = marks.get_indentation_score()
        variables_score = marks.get_variables_score()
        comment_score = marks.get_comment_score()
        total_score = marks.get_total_score()
        efficient_score = 0
        if execution_score > 4:
            efficient_score = 1
            execution_score -= 1
        self._feedback_document.tables[2].cell(
            (2 + int(execution_score)), 2).text = '{:g}'.format(
            execution_score)
        self._feedback_document.tables[2].cell(
            9, 2).text = str(indentation_score)
        self._feedback_document.tables[2].cell(
            10, 2).text = str(variables_score)
        self._feedback_document.tables[2].cell(
            11, 2).text = str(efficient_score)
        self._feedback_document.tables[2].cell(
            13 + int(comment_score), 2).text = str(comment_score)

        self._feedback_document.tables[2].cell(
            16, 2).text = '{:g}'.format(total_score)
        self._feedback_document.save(student_dir+'/../'+filename)

    def _write_student_name_to_document(
            self, student_dir, student_dir_name, student_name):
        """
        Internal method, write out student details to the feedback 
        document.
        """
        #default cell for student's firstname lastname
        filename = self._feedback_doc_name.replace(
            'username', student_dir_name)
        self._feedback_document.tables[0].cell(1,0).text = student_name
        self._feedback_document.tables[0].cell(1,1).text = student_dir_name
        self._feedback_document.tables[0].cell(1,2).text = self._marker_name
        self._feedback_document.save(student_dir+'/../'+filename)
        #print student_dir+'/'+filename

    def _write_comments_to_document(
            self, student_dir, student_dir_name, marks):
        """
        Internal method, write out marking comments to the feedback 
        document.
        """
        filename = os.path.basename(
            self._feedback_doc_name).replace('username', student_dir_name)

        feedback_document = Document(student_dir+'/../'+filename)
        feedback_document.add_heading('Program Comments', 2)

        feedback_document.add_heading('Program input', 3)
        inputs = marks.get_input().splitlines()

        for line in inputs:
            feedback_document.add_paragraph(line, style='CodeChunk')

        extra_program_inputs = marks.get_extra_program_inputs()
        for extra in extra_program_inputs:
            feedback_document.add_heading(extra[0], 3)
            extra_lines = extra[1].splitlines()
            for line in extra_lines:
                feedback_document.add_paragraph(line, style='CodeChunk')
    
        feedback_document.add_heading('Program output', 3)

        output = marks.get_output().splitlines()
        for line in output:
            feedback_document.add_paragraph(line, style='CodeChunk')

        extra_program_outputs = marks.get_extra_program_outputs()
        for extra in extra_program_outputs:
            feedback_document.add_heading(extra[0], 3)
            extra_lines = extra[1].splitlines()
            for line in extra_lines:
                feedback_document.add_paragraph(line, style='CodeChunk')

        feedback_document.add_heading('Execution comments', 3)

        comments = marks.get_execution_comments().splitlines()
        for line in comments:
            feedback_document.add_paragraph(line, style='CodeChunkComment')

        feedback_document.add_heading('Your code', 3)
        program = marks.get_full_program().splitlines()
        comments = marks.get_error_list()
        line_num = 1
        for line in program:
            line_encoded = line.decode('ascii', 'replace')
            highlight = False
            for comment in comments:
                if comment[0] == line_num:
                    highlight = True
            if highlight:
                feedback_document.add_paragraph(
                    str(line_num) + '\t: ' + \
                    line_encoded, style='CodeChunkHighlight')
                for comment in comments:
                    if comment[0] == line_num:
                        feedback_document.add_paragraph(
                            '\t  ' + comment[1], style='CodeChunkComment')
            else:
                feedback_document.add_paragraph(
                    str(line_num) + '\t: ' + line_encoded, style='CodeChunk')
            line_num += 1
        for comment in comments:
            if comment[0] == 0:
                feedback_document.add_paragraph(
                    '\t  ' + comment[1], style='CodeChunkComment')

        feedback_document.save(student_dir+'/../'+filename)

    def _construct_name_map(self):
        """
        Internal method, collect the usernames and student names from a 
        provided Excel document. The details are used to populate the
        feedback document with student details.
        """
        username_index = 0
        is_constructing_name_map = False

        for i in range(self._marking_sheet.nrows):
            if is_constructing_name_map:
                username =  self._marking_sheet.row_values(i)[username_index]
                firstname = self._marking_sheet.row_values(i)[username_index-1]
                lastname =  self._marking_sheet.row_values(i)[username_index-2]
                self._name_map[username]=[firstname, lastname]

            elif self._marking_sheet.row_values(i).count('Username') is 1:
                username_index = self._marking_sheet.row_values(i).index(
                    'Username')
                is_constructing_name_map = True


# Task number
# Folder containing students' folders
# Marker's initials (optional)
# Temp build folder (optional)
# Feedback template (optional)
# Student details list (optional)
# Summary output (optional)

# Run as a script, but skip this bit if the code is being imported
if __name__ == "__main__":
    start = time()
    parser = ArgumentParser(
        description='Batch marker for 4001COMP Java programming.')

    # Required parameters
    parser.add_argument(
        'task', metavar='TASK', type=int, help='Task number (e.g. 1)')
    parser.add_argument(
        'workdir', metavar='WORK', type=str, 
        help='Folder containing students\' folders (e.g. ./DLJ)')

    # Optional parameters
    parser.add_argument(
        '-i', '--initials', metavar='INITIALS', type=str, 
        help='Marker\'s initials (default Master)', default='Master')
    parser.add_argument(
        '-b', '--builddir', metavar='BUILD', type=str, 
        help='Folder to output build files to (default ./build)', 
        default='./build')
    parser.add_argument(
        '-t', '--template', metavar='TEMPLATE', type=str, 
        help='Word feedback sheeet template (default ./feedback_username.docx)', 
        default='./feedback_username.docx')
    parser.add_argument(
        '-d', '--details', metavar='DETAILS', type=str, 
        help='Excel student list (default ./4001COMP Marking 2014-15.xlsx)', 
        default='./4001COMP Marking 2014-15.xlsx')

    parser.add_argument(
        '-s', '--summary', metavar='SUMMARY', type=str, 
        help='Summary of marks as a CSV file (default ./summary.csv)', 
        default='./summary.csv')

    # Apply these arguments
    args = parser.parse_args()
    batchmark = BatchMark(
        args.task, args.workdir, args.initials, args.builddir, args.template, 
        args.details, args.summary)
    batchmark.go()

    end = time()
    duration = end - start
    print 'Time taken: {}'.format(duration)

